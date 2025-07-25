#!/usr/bin/env python3
"""
Enhanced PDF to Word Converter - Perfect layout preservation using pdf2docx
Converts .pdf files to .docx with exact formatting and layout matching
No AI services - local processing only
"""

import os
import sys
import json
import uuid
import time
import tempfile
import logging
from pathlib import Path
from typing import Dict, Any, Optional, Tuple

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

def log_message(message):
    """Log message with timestamp"""
    timestamp = time.strftime("%Y-%m-%d %H:%M:%S")
    print(f"[{timestamp}] {message}")

def detect_pdf_type(pdf_path):
    """Detect if PDF is text-based or scanned (image-based)"""
    try:
        import fitz  # PyMuPDF
        
        doc = fitz.open(pdf_path)
        total_chars = 0
        total_pages = len(doc)
        
        # Sample first few pages to determine type
        sample_pages = min(3, total_pages)
        
        for page_num in range(sample_pages):
            page = doc[page_num]
            text = page.get_text()
            total_chars += len(text.strip())
        
        doc.close()
        
        # If less than 50 characters per page on average, likely scanned
        avg_chars_per_page = total_chars / sample_pages if sample_pages > 0 else 0
        is_scanned = avg_chars_per_page < 50
        
        log_message(f"PDF analysis: {total_pages} pages, {avg_chars_per_page:.1f} chars/page avg, {'scanned' if is_scanned else 'text-based'}")
        
        return is_scanned, total_pages, avg_chars_per_page
        
    except Exception as e:
        log_message(f"PDF type detection failed: {str(e)}")
        return False, 0, 0

def convert_pdf_to_docx_perfect(pdf_path, output_dir):
    """Convert PDF to DOCX using pdf2docx with perfect layout preservation"""
    try:
        from pdf2docx import Converter
        
        # Generate output path
        input_stem = Path(pdf_path).stem
        docx_path = os.path.join(output_dir, f"{input_stem}.docx")
        
        log_message(f"Starting pdf2docx conversion: {pdf_path} -> {docx_path}")
        
        # Create converter and convert with optimal settings
        cv = Converter(pdf_path)
        
        # Convert with advanced settings for better layout preservation
        cv.convert(
            docx_path, 
            start=0, 
            end=None,
            multi_processing=True,
            cpu_count=2
        )
        cv.close()
        
        if os.path.exists(docx_path) and os.path.getsize(docx_path) > 0:
            log_message(f"pdf2docx conversion successful: {docx_path}")
            return True, docx_path, "Perfect layout conversion with pdf2docx"
        else:
            return False, None, "pdf2docx conversion failed - no output file generated"
            
    except ImportError:
        log_message("pdf2docx library not available, falling back to OCR method")
        return False, None, "pdf2docx library not installed"
    except Exception as e:
        log_message(f"pdf2docx conversion error: {str(e)}")
        return False, None, f"pdf2docx conversion failed: {str(e)}"

def convert_scanned_pdf_with_ocr(pdf_path, output_dir):
    """Convert scanned PDF using OCR with enhanced accuracy"""
    try:
        from pdf2image import convert_from_path
        import pytesseract
        from PIL import Image
        from docx import Document
        from docx.shared import Inches
        
        log_message("Converting scanned PDF with OCR...")
        
        # Convert PDF to images with high DPI for better OCR
        images = convert_from_path(pdf_path, dpi=300, fmt='png')
        
        # Create Word document
        doc = Document()
        
        input_stem = Path(pdf_path).stem
        docx_path = os.path.join(output_dir, f"{input_stem}.docx")
        
        for i, image in enumerate(images):
            log_message(f"Processing page {i+1} with OCR...")
            
            # Configure tesseract for better accuracy
            custom_config = r'--oem 3 --psm 6 -c tessedit_char_whitelist=0123456789ABCDEFGHIJKLMNOPQRSTUVWXYZabcdefghijklmnopqrstuvwxyz.,!?;:()\'"- '
            
            # Perform OCR on the image
            text = pytesseract.image_to_string(image, lang='eng', config=custom_config)
            
            # Add page heading
            if i > 0:
                doc.add_page_break()
            
            heading = doc.add_heading(f'Page {i+1}', level=2)
            
            # Process OCR text into paragraphs
            paragraphs = text.split('\n\n')
            for paragraph_text in paragraphs:
                paragraph_text = paragraph_text.strip()
                if paragraph_text and len(paragraph_text) > 3:
                    doc.add_paragraph(paragraph_text)
        
        # Save the document
        doc.save(docx_path)
        
        if os.path.exists(docx_path) and os.path.getsize(docx_path) > 0:
            log_message(f"OCR conversion successful: {docx_path}")
            return True, docx_path, f"OCR conversion completed for {len(images)} pages"
        else:
            return False, None, "OCR conversion failed - no output generated"
            
    except ImportError as e:
        log_message(f"OCR libraries not available: {str(e)}")
        return False, None, "OCR libraries not installed (pdf2image, pytesseract, python-docx)"
    except Exception as e:
        log_message(f"OCR conversion error: {str(e)}")
        return False, None, f"OCR conversion failed: {str(e)}"

def main():
    """Main conversion function"""
    if len(sys.argv) != 3:
        print(json.dumps({
            "success": False,
            "error": "Usage: python enhanced_pdf_converter.py <input_pdf> <output_dir>"
        }))
        sys.exit(1)
    
    input_path = sys.argv[1]
    output_dir = sys.argv[2]
    
    try:
        # Validate input file
        if not os.path.exists(input_path):
            print(json.dumps({
                "success": False,
                "error": "Input file does not exist"
            }))
            sys.exit(1)
        
        # Check file extension
        if not input_path.lower().endswith('.pdf'):
            print(json.dumps({
                "success": False,
                "error": "Only .pdf files are supported"
            }))
            sys.exit(1)
        
        # Check file size (50MB limit)
        file_size = os.path.getsize(input_path)
        if file_size > 50 * 1024 * 1024:
            print(json.dumps({
                "success": False,
                "error": "File size exceeds 50MB limit"
            }))
            sys.exit(1)
        
        log_message(f"Processing PDF: {input_path} ({file_size} bytes)")
        
        # Ensure output directory exists
        os.makedirs(output_dir, exist_ok=True)
        
        # Detect PDF type
        is_scanned, total_pages, avg_chars = detect_pdf_type(input_path)
        
        if total_pages == 0:
            print(json.dumps({
                "success": False,
                "error": "Could not read PDF file or file is empty"
            }))
            sys.exit(1)
        
        # Try pdf2docx first for perfect layout preservation
        log_message("Attempting high-quality conversion with pdf2docx...")
        success, docx_path, message = convert_pdf_to_docx_perfect(input_path, output_dir)
        
        if not success and is_scanned:
            log_message("pdf2docx failed, trying OCR for scanned PDF...")
            success, docx_path, message = convert_scanned_pdf_with_ocr(input_path, output_dir)
        
        if success:
            output_size = os.path.getsize(docx_path)
            
            print(json.dumps({
                "success": True,
                "output_path": docx_path,
                "file_size": output_size,
                "total_pages": total_pages,
                "conversion_method": "pdf2docx" if "pdf2docx" in message else "OCR",
                "message": message,
                "is_scanned": is_scanned
            }))
        else:
            print(json.dumps({
                "success": False,
                "error": message
            }))
            sys.exit(1)
    
    except Exception as e:
        print(json.dumps({
            "success": False,
            "error": f"Unexpected error: {str(e)}"
        }))
        sys.exit(1)

if __name__ == "__main__":
    main()