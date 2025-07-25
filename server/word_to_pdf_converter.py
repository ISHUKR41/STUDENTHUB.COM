#!/usr/bin/env python3
"""
Word to PDF Converter - Real conversion using LibreOffice
Converts .docx files to .pdf with accurate layout preservation
No AI services - local processing only
"""

import os
import sys
import json
import uuid
import time
import subprocess
import shutil
from pathlib import Path

def log_message(message):
    """Log message with timestamp"""
    timestamp = time.strftime("%Y-%m-%d %H:%M:%S")
    print(f"[{timestamp}] {message}")

def validate_docx_file(file_path):
    """Validate if file is a proper DOCX file"""
    try:
        from docx import Document
        doc = Document(file_path)
        # Try to access basic properties to ensure it's valid
        paragraphs = len(doc.paragraphs)
        return True, f"Valid DOCX with {paragraphs} paragraphs"
    except Exception as e:
        return False, f"Invalid DOCX file: {str(e)}"

def convert_docx_to_pdf_libreoffice(input_path, output_dir):
    """Convert DOCX to PDF using LibreOffice headless mode"""
    try:
        # Use LibreOffice for conversion (best layout preservation)
        cmd = [
            "libreoffice",
            "--headless",
            "--convert-to", "pdf",
            "--outdir", output_dir,
            input_path
        ]
        
        log_message(f"Running LibreOffice conversion: {' '.join(cmd)}")
        
        # Run conversion with timeout
        process = subprocess.run(
            cmd,
            capture_output=True,
            text=True,
            timeout=60  # 60 second timeout
        )
        
        if process.returncode == 0:
            # Find the generated PDF file
            input_stem = Path(input_path).stem
            pdf_path = os.path.join(output_dir, f"{input_stem}.pdf")
            
            if os.path.exists(pdf_path) and os.path.getsize(pdf_path) > 0:
                log_message(f"Conversion successful: {pdf_path}")
                return True, pdf_path, "LibreOffice conversion successful"
            else:
                return False, None, "PDF file was not generated or is empty"
        else:
            error_msg = process.stderr or process.stdout or "Unknown error"
            return False, None, f"LibreOffice conversion failed: {error_msg}"
            
    except subprocess.TimeoutExpired:
        return False, None, "Conversion timeout (60 seconds exceeded)"
    except FileNotFoundError:
        return False, None, "LibreOffice not found. Please install LibreOffice."
    except Exception as e:
        return False, None, f"Conversion error: {str(e)}"

def convert_docx_to_pdf_fallback(input_path, output_dir):
    """Fallback conversion using python-docx and reportlab"""
    try:
        from docx import Document
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.units import inch
        
        # Read DOCX content
        doc = Document(input_path)
        
        # Generate output PDF path
        input_stem = Path(input_path).stem
        pdf_path = os.path.join(output_dir, f"{input_stem}.pdf")
        
        # Create PDF
        pdf_doc = SimpleDocTemplate(pdf_path, pagesize=letter)
        styles = getSampleStyleSheet()
        story = []
        
        # Add content from DOCX
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                # Determine style based on paragraph formatting
                if paragraph.style.name.startswith('Heading'):
                    style = styles['Heading1']
                else:
                    style = styles['Normal']
                
                story.append(Paragraph(paragraph.text, style))
                story.append(Spacer(1, 0.2*inch))
        
        # Build PDF
        pdf_doc.build(story)
        
        if os.path.exists(pdf_path) and os.path.getsize(pdf_path) > 0:
            log_message(f"Fallback conversion successful: {pdf_path}")
            return True, pdf_path, "Fallback conversion successful"
        else:
            return False, None, "Fallback PDF generation failed"
            
    except ImportError:
        return False, None, "Required libraries not available for fallback conversion"
    except Exception as e:
        return False, None, f"Fallback conversion error: {str(e)}"

def main():
    """Main conversion function"""
    if len(sys.argv) != 3:
        print(json.dumps({
            "success": False,
            "error": "Usage: python word_to_pdf_converter.py <input_docx> <output_dir>"
        }))
        sys.exit(1)
    
    input_path = sys.argv[1]
    output_dir = sys.argv[2]
    
    try:
        # Validate input file
        if not os.path.exists(input_path):
            print(json.dumps({
                "success": False,
                "error": "Input file does not exist"
            }))
            sys.exit(1)
        
        # Check file extension
        if not input_path.lower().endswith('.docx'):
            print(json.dumps({
                "success": False,
                "error": "Only .docx files are supported"
            }))
            sys.exit(1)
        
        # Validate DOCX file
        is_valid, validation_msg = validate_docx_file(input_path)
        if not is_valid:
            print(json.dumps({
                "success": False,
                "error": validation_msg
            }))
            sys.exit(1)
        
        log_message(f"Validation passed: {validation_msg}")
        
        # Ensure output directory exists
        os.makedirs(output_dir, exist_ok=True)
        
        # Try LibreOffice conversion first
        log_message("Attempting LibreOffice conversion...")
        success, pdf_path, message = convert_docx_to_pdf_libreoffice(input_path, output_dir)
        
        if not success:
            log_message(f"LibreOffice failed: {message}")
            log_message("Trying fallback conversion...")
            success, pdf_path, message = convert_docx_to_pdf_fallback(input_path, output_dir)
        
        if success:
            # Get file stats
            file_size = os.path.getsize(pdf_path)
            
            print(json.dumps({
                "success": True,
                "output_path": pdf_path,
                "file_size": file_size,
                "message": message,
                "conversion_method": "libreoffice" if "LibreOffice" in message else "fallback"
            }))
        else:
            print(json.dumps({
                "success": False,
                "error": message
            }))
            sys.exit(1)
    
    except Exception as e:
        print(json.dumps({
            "success": False,
            "error": f"Unexpected error: {str(e)}"
        }))
        sys.exit(1)

if __name__ == "__main__":
    main()