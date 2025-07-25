import os
import uuid
import time
import hashlib
import logging
from pathlib import Path
from typing import Dict, Any, Optional, Tuple
import fitz  # PyMuPDF
import pytesseract
from pdf2image import convert_from_path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image
import tempfile
import shutil

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class PDFToWordConverter:
    def __init__(self, temp_dir: Optional[str] = None):
        self.temp_dir = temp_dir or tempfile.gettempdir()
        self.active_files = {}  # Track files with expiry times
        
    def validate_pdf(self, file_path: str, max_size_mb: int = 50) -> Dict[str, Any]:
        """Validate PDF file"""
        try:
            # Check file exists
            if not os.path.exists(file_path):
                return {"valid": False, "error": "File not found"}
            
            # Check file size
            file_size = os.path.getsize(file_path)
            if file_size > max_size_mb * 1024 * 1024:
                return {"valid": False, "error": f"File too large. Max size: {max_size_mb}MB"}
            
            # Check if it's a valid PDF
            try:
                doc = fitz.open(file_path)
                page_count = len(doc)
                doc.close()
                
                if page_count == 0:
                    return {"valid": False, "error": "PDF has no pages"}
                
                return {"valid": True, "pages": page_count, "size_mb": round(file_size / (1024 * 1024), 2)}
                
            except Exception as e:
                return {"valid": False, "error": "Invalid or corrupted PDF file"}
                
        except Exception as e:
            logger.error(f"PDF validation error: {str(e)}")
            return {"valid": False, "error": "File validation failed"}
    
    def detect_pdf_type(self, file_path: str) -> Dict[str, Any]:
        """Detect if PDF is text-based or scanned (image-based)"""
        try:
            doc = fitz.open(file_path)
            total_chars = 0
            total_pages = len(doc)
            
            # Sample first 5 pages or all pages if fewer
            sample_pages = min(5, total_pages)
            
            for page_num in range(sample_pages):
                page = doc.load_page(page_num)
                text = page.get_text() if hasattr(page, 'get_text') else ""
                total_chars += len(text.strip())
            
            doc.close()
            
            # Calculate character density
            char_density = total_chars / sample_pages if sample_pages > 0 else 0
            
            # If average characters per page is less than 100, likely scanned
            is_scanned = char_density < 100
            
            return {
                "is_scanned": is_scanned,
                "char_density": char_density,
                "total_pages": total_pages,
                "sample_pages": sample_pages
            }
            
        except Exception as e:
            logger.error(f"PDF type detection error: {str(e)}")
            return {"error": "Could not analyze PDF type"}
    
    def extract_text_from_pdf(self, file_path: str) -> Dict[str, Any]:
        """Extract text from text-based PDF"""
        try:
            doc = fitz.open(file_path)
            text_content = []
            
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                text = page.get_text() if hasattr(page, 'get_text') else ""
                
                # Get text with formatting info
                blocks = page.get_text("dict") if hasattr(page, 'get_text') else {}
                page_content = {
                    "page_number": page_num + 1,
                    "text": text,
                    "blocks": blocks
                }
                text_content.append(page_content)
            
            doc.close()
            
            return {"success": True, "content": text_content}
            
        except Exception as e:
            logger.error(f"Text extraction error: {str(e)}")
            return {"success": False, "error": "Failed to extract text from PDF"}
    
    def try_enhanced_conversion(self, pdf_path: str) -> Dict[str, Any]:
        """Try enhanced conversion using pdf2docx for perfect layout preservation"""
        try:
            import subprocess
            import json
            
            # Call the enhanced converter script
            result = subprocess.run([
                'python3',
                os.path.join(os.path.dirname(__file__), 'enhanced_pdf_converter.py'),
                pdf_path,
                self.temp_dir
            ], capture_output=True, text=True, timeout=300)  # 5 minute timeout
            
            if result.returncode == 0:
                output = json.loads(result.stdout.strip())
                if output.get("success"):
                    logger.info(f"Enhanced conversion successful: {output.get('message')}")
                    return {
                        "success": True,
                        "enhanced_conversion": True,
                        "output_path": output.get("output_path"),
                        "conversion_method": output.get("conversion_method"),
                        "message": output.get("message")
                    }
                else:
                    logger.warning(f"Enhanced conversion failed: {output.get('error')}")
                    return {"success": False, "error": output.get("error")}
            else:
                logger.warning(f"Enhanced converter process failed with code {result.returncode}")
                return {"success": False, "error": "Enhanced converter process failed"}
                
        except Exception as e:
            logger.warning(f"Enhanced conversion error: {str(e)}")
            return {"success": False, "error": f"Enhanced conversion failed: {str(e)}"}
    
    def extract_text_with_ocr(self, file_path: str) -> Dict[str, Any]:
        """Extract text from scanned PDF using OCR"""
        try:
            # Convert PDF to images
            images = convert_from_path(file_path, dpi=300)
            text_content = []
            
            for i, image in enumerate(images):
                # Perform OCR on each page
                try:
                    text = pytesseract.image_to_string(image, config='--psm 6')
                    page_content = {
                        "page_number": i + 1,
                        "text": text,
                        "method": "ocr"
                    }
                    text_content.append(page_content)
                except Exception as ocr_error:
                    logger.warning(f"OCR failed for page {i + 1}: {str(ocr_error)}")
                    text_content.append({
                        "page_number": i + 1,
                        "text": "",
                        "error": "OCR failed for this page"
                    })
            
            return {"success": True, "content": text_content}
            
        except Exception as e:
            logger.error(f"OCR extraction error: {str(e)}")
            return {"success": False, "error": "OCR processing failed"}
    
    def create_word_document(self, content: list, output_path: str) -> Dict[str, Any]:
        """Create Word document from extracted content"""
        try:
            doc = Document()
            
            # Set document margins
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1)
                section.right_margin = Inches(1)
            
            for page_data in content:
                page_num = page_data.get("page_number", 1)
                text = page_data.get("text", "")
                
                if text.strip():
                    # Add page break for pages after the first
                    if page_num > 1:
                        doc.add_page_break()
                    
                    # Add page header
                    if page_num > 1:
                        header = doc.add_paragraph(f"Page {page_num}")
                        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        header_format = header.runs[0].font
                        header_format.size = Pt(10)
                        header_format.italic = True
                    
                    # Process text blocks
                    paragraphs = text.split('\n\n')
                    for para_text in paragraphs:
                        if para_text.strip():
                            para = doc.add_paragraph(para_text.strip())
                            # Set font
                            for run in para.runs:
                                run.font.name = 'Calibri'
                                run.font.size = Pt(11)
            
            # Save document
            doc.save(output_path)
            
            return {"success": True, "output_path": output_path}
            
        except Exception as e:
            logger.error(f"Word document creation error: {str(e)}")
            return {"success": False, "error": "Failed to create Word document"}
    
    def generate_secure_download_id(self, file_path: str) -> str:
        """Generate secure download ID with timestamp"""
        timestamp = str(int(time.time()))
        file_info = f"{file_path}_{timestamp}_{uuid.uuid4()}"
        return hashlib.sha256(file_info.encode()).hexdigest()
    
    def convert_pdf_to_word(self, pdf_path: str, client_ip: Optional[str] = None) -> Dict[str, Any]:
        """Main conversion function"""
        try:
            logger.info(f"Starting PDF conversion: {pdf_path} from IP: {client_ip}")
            
            # Validate PDF
            validation = self.validate_pdf(pdf_path)
            if not validation["valid"]:
                return {"success": False, "error": validation["error"]}
            
            # Detect PDF type
            pdf_type = self.detect_pdf_type(pdf_path)
            if "error" in pdf_type:
                return {"success": False, "error": pdf_type["error"]}
            
            # Try enhanced converter first for perfect layout preservation
            enhanced_result = self.try_enhanced_conversion(pdf_path)
            if enhanced_result["success"]:
                logger.info("Using enhanced pdf2docx conversion")
                extraction_result = enhanced_result
            else:
                # Fallback to existing logic
                if pdf_type["is_scanned"]:
                    logger.info("Processing scanned PDF with OCR")
                    extraction_result = self.extract_text_with_ocr(pdf_path)
                else:
                    logger.info("Processing text-based PDF")
                    extraction_result = self.extract_text_from_pdf(pdf_path)
            
            if not extraction_result["success"]:
                return {"success": False, "error": extraction_result["error"]}
            
            # If enhanced conversion was successful, use the generated file directly
            if extraction_result.get("enhanced_conversion"):
                output_path = extraction_result["output_path"]
                word_result = {"success": True, "output_path": output_path}
            else:
                # Generate output file path
                output_filename = f"converted_{uuid.uuid4()}.docx"
                output_path = os.path.join(self.temp_dir, output_filename)
                
                # Create Word document using fallback method
                word_result = self.create_word_document(extraction_result["content"], output_path)
            if not word_result["success"]:
                return {"success": False, "error": word_result["error"]}
            
            # Generate secure download ID
            download_id = self.generate_secure_download_id(output_path)
            expiry_time = time.time() + (4 * 60)  # 4 minutes from now
            
            # Store file info for cleanup
            self.active_files[download_id] = {
                "file_path": output_path,
                "pdf_path": pdf_path,
                "expiry_time": expiry_time,
                "created_at": time.time(),
                "client_ip": client_ip
            }
            
            # Get file size
            file_size = os.path.getsize(output_path)
            
            logger.info(f"PDF conversion completed. Download ID: {download_id}")
            
            return {
                "success": True,
                "download_id": download_id,
                "expiry_time": expiry_time,
                "file_size": file_size,
                "pdf_type": "scanned" if pdf_type["is_scanned"] else "text-based",
                "pages": validation["pages"]
            }
            
        except Exception as e:
            logger.error(f"Conversion error: {str(e)}")
            return {"success": False, "error": "Conversion failed"}
    
    def get_download_file(self, download_id: str) -> Optional[Tuple[str, Dict[str, Any]]]:
        """Get file path for download"""
        if download_id not in self.active_files:
            return None
        
        file_info = self.active_files[download_id]
        
        # Check if file has expired
        if time.time() > file_info["expiry_time"]:
            self.cleanup_file(download_id)
            return None
        
        # Check if file still exists
        if not os.path.exists(file_info["file_path"]):
            self.cleanup_file(download_id)
            return None
        
        return file_info["file_path"], file_info
    
    def cleanup_file(self, download_id: str):
        """Clean up files for a specific download ID"""
        if download_id in self.active_files:
            file_info = self.active_files[download_id]
            
            # Remove Word file
            if os.path.exists(file_info["file_path"]):
                try:
                    os.remove(file_info["file_path"])
                    logger.info(f"Cleaned up Word file: {file_info['file_path']}")
                except Exception as e:
                    logger.error(f"Failed to remove Word file: {e}")
            
            # Remove PDF file
            if os.path.exists(file_info["pdf_path"]):
                try:
                    os.remove(file_info["pdf_path"])
                    logger.info(f"Cleaned up PDF file: {file_info['pdf_path']}")
                except Exception as e:
                    logger.error(f"Failed to remove PDF file: {e}")
            
            # Remove from tracking
            del self.active_files[download_id]
    
    def cleanup_expired_files(self):
        """Clean up all expired files"""
        current_time = time.time()
        expired_ids = []
        
        for download_id, file_info in self.active_files.items():
            if current_time > file_info["expiry_time"]:
                expired_ids.append(download_id)
        
        for download_id in expired_ids:
            self.cleanup_file(download_id)
        
        logger.info(f"Cleaned up {len(expired_ids)} expired files")